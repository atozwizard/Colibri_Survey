#!/usr/bin/env python3
"""기업용 pandoc reference.docx 테마 생성.
- 본문/제목/헤딩 폰트·색상(네이비 계열), 한글 eastAsia 폰트 지정.
- 상단 머리글(보고서명), 하단 바닥글(페이지 번호).
pandoc --reference-doc=report/reference.docx 로 사용.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "report" / "reference.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
GRAY = RGBColor(0x40, 0x40, 0x40)
LATIN = "Helvetica Neue"
KO = "Apple SD Gothic Neo"   # macOS 한글; Word는 없으면 Malgun Gothic 등으로 대체


def set_fonts(style, latin=LATIN, ko=KO, size=None, color=None, bold=None):
    f = style.font
    f.name = latin
    if size is not None:
        f.size = Pt(size)
    if color is not None:
        f.color.rgb = color
    if bold is not None:
        f.bold = bold
    # eastAsia(한글) 폰트 지정
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), ko)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def add_page_number(paragraph):
    """바닥글에 'PAGE' 필드 삽입."""
    run = paragraph.add_run()
    fld = run._r.makeelement(qn("w:fldSimple"), {qn("w:instr"): "PAGE"})
    run._r.append(fld)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    d = Document()

    st = d.styles
    set_fonts(st["Normal"], size=10.5, color=RGBColor(0x22, 0x22, 0x22))
    set_fonts(st["Heading 1"], size=18, color=NAVY, bold=True)
    set_fonts(st["Heading 2"], size=14, color=BLUE, bold=True)
    set_fonts(st["Heading 3"], size=12, color=GRAY, bold=True)
    for name in ("Title", "Subtitle"):
        if name in [s.name for s in st]:
            pass
    try:
        set_fonts(st["Title"], size=28, color=NAVY, bold=True)
    except KeyError:
        pass
    try:
        set_fonts(st["Subtitle"], size=14, color=GRAY, bold=False)
    except KeyError:
        pass
    # 코드/표 가독성
    for name in ("Source Code", "Verbatim Char", "macro"):
        try:
            set_fonts(st[name], latin="Menlo", ko=KO, size=9)
        except KeyError:
            pass

    sec = d.sections[0]
    # 머리글
    h = sec.header.paragraphs[0]
    h.text = "colibrì 종합 서베이 보고서 · Colibri_Survey"
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
    # 바닥글(페이지 번호)
    f = sec.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run("— ")
    add_page_number(f)
    f.add_run(" —")
    for r in f.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY

    d.save(str(OUT))
    print(f"reference.docx 작성: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
